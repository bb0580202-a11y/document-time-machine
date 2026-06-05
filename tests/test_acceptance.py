"""AGENT_SPEC §7 验收：1 智能忽略 2 基本快照 3 反悔 4 安全还原 5 分支
6 搬家找回 7 膨胀 8 离线 9 完整性 10 备注。"""
from pathlib import Path

from docx import Document

from dtm.engine.repo import GitRepo
from dtm.engine import backup, identity, restore, meta, stats
from tests.conftest import git_log_files


def _setup(folder):
    repo = GitRepo(folder)
    repo.init()
    identity.write_identity(folder)
    (folder / ".gitignore").write_text(backup.gitignore_text())
    return repo


def _save(folder, name, text):
    d = Document()
    d.add_paragraph(text)
    d.save(folder / name)


def test_s1_smart_ignore(folder):
    repo = _setup(folder)
    _save(folder, "正文.docx", "A")
    (folder / "~$正文.docx").write_bytes(b"lock")
    (folder / "temp.tmp").write_text("junk")
    big = folder / "大图.bin"
    big.write_bytes(b"\0" * (51 * 1024 * 1024))
    res = backup.do_backup(repo, folder)
    files = git_log_files(folder)
    assert "正文.docx" in files
    assert "~$正文.docx" not in files and "temp.tmp" not in files
    assert "大图.bin" in files and any("大图" in w for w in res.warnings)


def test_s2_basic_snapshots(folder):
    repo = _setup(folder)
    for t in ("A", "B", "C"):
        _save(folder, "正文.docx", t)
        backup.do_backup(repo, folder)
    assert len(repo.log()) == 3


def test_s3_undo_after_restore_keeps_both(folder):
    repo = _setup(folder)
    _save(folder, "正文.docx", "第三段=内容A")
    cidA = backup.do_backup(repo, folder).commit_id
    _save(folder, "正文.docx", "第三段=内容B")
    backup.do_backup(repo, folder)
    res = restore.safe_restore(repo, folder, cidA, "正文.docx")
    # 能取回内容A
    assert "内容A" in Document(res.restored_path).paragraphs[0].text
    # 内容B 版本仍在历史中（INV-1）：从最新版本读回仍是内容B
    head = repo.log()[0].commit_id
    check_b = folder / "_checkB.docx"
    check_b.write_bytes(repo.show_file(head, "正文.docx"))
    assert "内容B" in Document(check_b).paragraphs[0].text


def test_s4_safe_restore_no_overwrite(folder):
    repo = _setup(folder)
    _save(folder, "正文.docx", "A")
    cidA = backup.do_backup(repo, folder).commit_id
    _save(folder, "正文.docx", "未保存改动")
    res = restore.safe_restore(repo, folder, cidA, "正文.docx")
    assert "未保存改动" in Document(folder / "正文.docx").paragraphs[0].text
    assert Path(res.restored_path) != folder / "正文.docx"


def test_s5_branch(folder):
    repo = _setup(folder)
    _save(folder, "正文.docx", "主线")
    base = backup.do_backup(repo, folder).commit_id
    repo.create_branch("路线-2", base)
    repo.checkout("路线-2")
    _save(folder, "正文.docx", "岔路")
    backup.do_backup(repo, folder)
    assert len(repo.log()) >= 2


def test_s6_relocate_after_move(folder, tmp_path):
    repo = _setup(folder)
    for t in ("A", "B", "C"):
        _save(folder, "正文.docx", t)
        backup.do_backup(repo, folder)
    ident = identity.read_identity(folder)
    moved = tmp_path / "另一个盘" / "论文改名"
    moved.parent.mkdir()
    folder.rename(moved)
    found = identity.find_repo_by_uuid(ident["uuid"], [tmp_path])
    assert found == moved
    assert len(GitRepo(moved).log()) == 3   # 历史完整


def test_s7_bloat_numbers(folder):
    repo = _setup(folder)
    for i in range(10):  # 验收里跑 10 次证明可测量；详尽 50 次见 experiments/RESULTS.md
        _save(folder, "正文.docx", f"内容 {i} " * 200)
        backup.do_backup(repo, folder)
    before = stats.repo_stats(repo).git_bytes
    repo.gc()
    after = stats.repo_stats(repo).git_bytes
    assert before > 0 and after > 0


def test_s8_offline_no_remote(folder):
    import subprocess
    _setup(folder)
    remotes = subprocess.run(["git", "-C", str(folder), "remote"],
                             capture_output=True, text=True).stdout.strip()
    assert remotes == ""


def test_s9_integrity_detects_corruption(folder):
    from dtm.engine.integrity import check
    _setup(folder)
    bad = folder / "坏.docx"
    bad.write_bytes(b"not a zip")
    ok, reason = check(bad)
    assert ok is False


def test_s10_note_no_new_version(folder):
    repo = _setup(folder)
    _save(folder, "正文.docx", "A")
    cid = backup.do_backup(repo, folder).commit_id
    before = len(repo.log())
    meta.set_note(repo, cid, "投稿前的版本")
    assert meta.get_note(repo, cid) == "投稿前的版本"
    assert len(repo.log()) == before          # 备注不产生新快照
    meta.set_note(repo, cid, "")
    assert meta.get_note(repo, cid) == ""
    # 被守护文件内容不受影响
    assert "A" in Document(folder / "正文.docx").paragraphs[0].text
