import json
from pathlib import Path
from docx import Document
from dtm.engine.repo import GitRepo
from dtm.engine import backup, identity, meta, registry
from dtm.app.bridge import Bridge


def _guarded(folder):
    repo = GitRepo(folder); repo.init(); identity.write_identity(folder)
    (folder / ".gitignore").write_text(backup.gitignore_text())
    Document().save(folder / "正文.docx")
    backup.do_backup(repo, folder)
    return repo


def test_head_returns_current_commit(folder):
    repo = _guarded(folder)
    assert len(repo.head()) == 40
    assert repo.head() == repo.log()[0].commit_id


def test_get_treemap_marks_real_head_as_current(folder):
    repo = _guarded(folder)
    Document().save(folder / "正文.docx"); backup.do_backup(repo, folder)  # 第二版=HEAD
    tm = Bridge(store=folder / "ignored.json").get_treemap(str(folder))
    json.dumps(tm)                                  # 纯 JSON
    current = [n for n in tm["nodes"] if n["is_current"]]
    assert len(current) == 1
    assert current[0]["full_id"] == repo.head()     # CP-3:当前=真 HEAD


def test_get_album_cards_have_card_fields(folder):
    repo = _guarded(folder)
    cid = repo.head(); meta.set_note(repo, cid, "导师说结论太弱")
    cards = Bridge(store=folder / "ignored.json").get_album(str(folder))
    assert cards[0]["title"] == "导师说结论太弱"      # 有备注→备注当标题
    assert "abs_minute" in cards[0] and "relative" in cards[0]
    assert cards[0]["delta_sign"] in ("up", "down", "flat")
    assert cards[0]["files"][0]["name"] == "正文.docx"      # 逐文件列(design §5.2)
    assert cards[0]["files"][0]["delta_sign"] == "up"        # 首版从无到有


def test_add_then_remove_folder(folder, tmp_path):
    store = tmp_path / "folders.json"
    Document().save(folder / "正文.docx")
    b = Bridge(store=store, search_roots=[tmp_path])
    entry = b.add_folder(str(folder))               # 未守护→init+登记
    assert entry["status"] == "active"
    assert any(f["uuid"] == entry["uuid"] for f in b.list_folders())
    b.remove_folder(entry["uuid"])
    assert b.list_folders() == []


def test_list_folders_persists_relocated_path(folder, tmp_path):
    store = tmp_path / "folders.json"
    Document().save(folder / "正文.docx")
    b = Bridge(store=store, search_roots=[tmp_path])
    entry = b.add_folder(str(folder))
    moved = folder.parent / "renamed_thesis"
    folder.rename(moved)
    out = b.list_folders()                           # GUI 侧:resolve + 持久化
    assert Path(out[0]["path"]) == moved.resolve()   # 返回找回的新路径
    # 且落回磁盘:重新 load 看到的是新路径(GUI 是唯一写者,负责回写)
    persisted = registry.load(store)
    assert Path(persisted[0].path) == moved.resolve()


def test_restore_takes_lock_and_writes_beside(folder, tmp_path):
    repo = _guarded(folder)
    cid = repo.head()
    b = Bridge(store=tmp_path / "s.json")
    out = b.restore(str(folder), cid, "正文.docx")
    assert Path(out["restored_path"]).exists()       # 旁边另存
    assert (folder / ".git" / "dtm.lock").exists() is False   # 用完即释放


def test_restore_version_restores_all_changed_files(folder, tmp_path):
    repo = _guarded(folder)                       # 首版:正文.docx
    big = Document()
    for _ in range(50):
        big.add_paragraph("正文内容" * 20)
    big.save(folder / "正文.docx")                # 正文变大
    Document().save(folder / "参考文献.docx")      # 新增第二个文件
    backup.do_backup(repo, folder)
    cid = repo.head()                             # 这一版动了 2 个文件
    b = Bridge(store=tmp_path / "s.json")
    out = b.restore_version(str(folder), cid)
    assert len(out["restored_paths"]) == 2        # 两个变更文件都另存到旁边
    for p in out["restored_paths"]:
        assert Path(p).exists()
        assert "_恢复自" in Path(p).name           # 旁边另存,不覆盖(INV-3)
    assert (folder / ".git" / "dtm.lock").exists() is False   # 用完即释放


def test_set_note_no_new_version(folder, tmp_path):
    repo = _guarded(folder)
    before = len(repo.log())
    b = Bridge(store=tmp_path / "s.json")
    b.set_note(str(folder), repo.head(), "投稿前别动")
    assert len(repo.log()) == before                 # 备注不产生新快照
    assert b.get_album(str(folder))[0]["note"] == "投稿前别动"


def test_peek_fingerprint_changes_on_new_version(folder, tmp_path):
    repo = _guarded(folder)
    b = Bridge(store=tmp_path / "s.json")
    fp1 = b.peek(str(folder))                         # GUI 轮询用的轻量指纹
    assert len(fp1["head"]) == 40 and fp1["count"] >= 1
    big = Document()
    for _ in range(30):
        big.add_paragraph("内容" * 40)
    big.save(folder / "正文.docx")
    backup.do_backup(repo, folder)                    # 后台又存了一版
    fp2 = b.peek(str(folder))
    assert fp2 != fp1                                 # 指纹变 → 窗口该自动刷新
    assert fp2["count"] == fp1["count"] + 1


def test_bridge_humanizes_unexpected_error(tmp_path, monkeypatch):
    # 非 DtmError 的技术异常(git/OSError/KeyError…)绝不能原样漏到用户(INV-5/6)
    from dtm.app import bridge as bridge_mod
    b = Bridge(store=tmp_path / "s.json")

    def boom(*a, **k):
        raise KeyError("internal_tech_detail_xyz")
    monkeypatch.setattr(bridge_mod.registry, "resolve", boom)
    try:
        b.list_folders()
        assert False, "应当抛出异常"
    except Exception as e:
        msg = str(e)
        assert "internal_tech_detail_xyz" not in msg   # 技术细节不泄露
        assert "KeyError" not in msg
        assert ("重启" in msg or "再试" in msg)          # 是人话兜底


def test_bridge_passes_through_human_dterror(tmp_path, monkeypatch):
    # 已是人话的 DtmError 原样上抛,不被兜底文案盖掉(保留具体有用信息)
    from dtm.app import bridge as bridge_mod
    from dtm.engine.errors import LockBusyError
    b = Bridge(store=tmp_path / "s.json")

    def busy(*a, **k):
        raise LockBusyError("另一个备份正在进行，请稍候再试。")
    monkeypatch.setattr(bridge_mod.registry, "resolve", busy)
    try:
        b.list_folders()
        assert False, "应当抛出异常"
    except Exception as e:
        assert "另一个备份正在进行" in str(e)            # 人话原样保留


def test_reveal_path_opens_dir_and_reveals_file(folder, tmp_path, monkeypatch):
    # A2/A3 共用:目录→open 打开文件夹;文件→open -R 在访达里选中
    from dtm.app import bridge as bridge_mod
    b = Bridge(store=tmp_path / "s.json")
    calls = []
    monkeypatch.setattr(bridge_mod.sys, "platform", "darwin")
    monkeypatch.setattr(bridge_mod.subprocess, "run", lambda args, **k: calls.append(args))
    b.reveal_path(str(folder))
    assert calls[-1] == ["open", str(folder)]
    f = folder / "正文.docx"; f.write_text("x")
    b.reveal_path(str(f))
    assert calls[-1] == ["open", "-R", str(f)]


def test_reveal_path_missing_raises_human(tmp_path):
    b = Bridge(store=tmp_path / "s.json")
    try:
        b.reveal_path(str(tmp_path / "nope"))
        assert False, "应当抛出异常"
    except Exception as e:
        assert "找不到" in str(e)                         # 人话(DtmError 原样)，不漏技术细节


def test_bridge_autostart_roundtrip(tmp_path, monkeypatch):
    # 注入 tmp backend(写 tmp)→ 不碰真实 ~/Library;纯文件操作不跑 launchctl
    from dtm.engine.autostart import MacAutoStart
    import dtm.app.bridge as bridge_mod
    backend = MacAutoStart(label="com.x.daemon", plist_dir=tmp_path)
    monkeypatch.setattr(bridge_mod, "_autostart_backend", lambda: backend)
    b = Bridge()
    assert b.get_autostart() is False
    b.set_autostart(True);  assert b.get_autostart() is True    # 写了 plist
    b.set_autostart(False); assert b.get_autostart() is False   # 删了 plist


def test_archive_sets_flag(folder, tmp_path):
    store = tmp_path / "folders.json"
    _guarded(folder)
    registry.add(store, folder)
    uuid = registry.load(store)[0].uuid
    b = Bridge(store=store, search_roots=[folder.parent])
    b.archive_folder(uuid)
    assert registry.load(store)[0].archived is True


def test_resume_backs_up_then_unsets(folder, tmp_path):
    store = tmp_path / "folders.json"
    repo = _guarded(folder)
    registry.add(store, folder)
    uuid = registry.load(store)[0].uuid
    b = Bridge(store=store, search_roots=[folder.parent])
    b.archive_folder(uuid)
    (folder / "正文.docx").write_bytes(b"changed during archive")  # 归档期改(无保护)
    before = len(repo.log())
    b.resume_folder(uuid)
    assert len(repo.log()) == before + 1            # resume 先补拍一版基线
    assert registry.load(store)[0].archived is False  # 再 archived=false
