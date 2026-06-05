"""某版改动文件 + 每文件大小增减方向(▲▼)。
核心陷阱:docx 是二进制,git diff --numstat 会返回 '-' 让 ▲▼ 坏掉;
故靠 git cat-file -s 比对本版/父版 blob 大小算方向。"""
from docx import Document
from dtm.engine.repo import GitRepo
from dtm.engine import backup, identity, listing


def _init(folder):
    repo = GitRepo(folder); repo.init(); identity.write_identity(folder)
    (folder / ".gitignore").write_text(backup.gitignore_text())
    return repo


def _big_doc(path, n):
    d = Document()
    for _ in range(n):
        d.add_paragraph("段落内容" * 20)
    d.save(path)


def test_first_version_all_up_and_hides_gitignore(folder):
    repo = _init(folder)
    Document().save(folder / "正文.docx")
    backup.do_backup(repo, folder)
    cf = listing.changed_files(repo, repo.head())
    names = [c.name for c in cf]
    assert "正文.docx" in names
    assert ".gitignore" not in names                 # dtm 自有簿记不进用户视图
    assert all(c.delta_sign == "up" for c in cf)     # 首版从无到有


def test_grow_then_shrink(folder):
    repo = _init(folder)
    Document().save(folder / "正文.docx"); backup.do_backup(repo, folder)
    _big_doc(folder / "正文.docx", 300); backup.do_backup(repo, folder)   # 变大
    grew = [c for c in listing.changed_files(repo, repo.head()) if c.name == "正文.docx"][0]
    assert grew.delta_sign == "up"
    Document().save(folder / "正文.docx"); backup.do_backup(repo, folder)  # 清空变小
    shrank = [c for c in listing.changed_files(repo, repo.head()) if c.name == "正文.docx"][0]
    assert shrank.delta_sign == "down"


def test_binary_docx_delta_not_broken(folder):
    repo = _init(folder)
    Document().save(folder / "正文.docx"); backup.do_backup(repo, folder)
    _big_doc(folder / "正文.docx", 500); backup.do_backup(repo, folder)
    sign = [c for c in listing.changed_files(repo, repo.head()) if c.name == "正文.docx"][0].delta_sign
    assert sign == "up"                              # 不是 None/坏(numstat 会坏)


def test_multi_file_each_own_sign(folder):
    repo = _init(folder)
    Document().save(folder / "a.docx")
    _big_doc(folder / "b.docx", 300)
    backup.do_backup(repo, folder)
    # 第二版:a 变大、b 清空变小
    _big_doc(folder / "a.docx", 300)
    Document().save(folder / "b.docx")
    backup.do_backup(repo, folder)
    cf = {c.name: c.delta_sign for c in listing.changed_files(repo, repo.head())}
    assert cf["a.docx"] == "up"
    assert cf["b.docx"] == "down"
