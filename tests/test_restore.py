from pathlib import Path

import pytest
from docx import Document

from dtm.engine.repo import GitRepo
from dtm.engine.errors import DtmError
from dtm.engine import backup, restore


def _commit_docx(folder, name, text):
    repo = GitRepo(folder)
    if not repo.is_repo():
        repo.init()
        (folder / ".gitignore").write_text(backup.gitignore_text())
    p = folder / name
    doc = Document()
    doc.add_paragraph(text)
    doc.save(p)
    return repo, backup.do_backup(repo, folder).commit_id


def test_restore_is_byte_identical(folder):
    """策略 A 零格式风险的硬证据：还原文件与当初保存的 docx 逐字节一致。"""
    repo = GitRepo(folder)
    repo.init()
    (folder / ".gitignore").write_text(backup.gitignore_text())
    p = folder / "正文.docx"
    doc = Document()
    doc.add_heading("某精算问题", level=0)
    for i in range(20):
        doc.add_paragraph(f"第{i}段：含公式与样式的正文内容。")
    doc.save(p)
    original_bytes = p.read_bytes()
    cid = backup.do_backup(repo, folder).commit_id
    # 把当前文件改掉，逼还原必须从历史取回
    doc2 = Document()
    doc2.add_paragraph("完全不同的内容")
    doc2.save(p)
    res = restore.safe_restore(repo, folder, cid, "正文.docx")
    assert Path(res.restored_path).read_bytes() == original_bytes  # 一个 bit 都不差


def test_restore_name_has_second_precision_timestamp(folder):
    """还原副本名带秒级时间戳(唯一、自解释):正文_恢复自YYYYMMDD-HHMMSS.docx。"""
    import re
    repo, cid = _commit_docx(folder, "正文.docx", "内容")
    out = Path(restore.safe_restore(repo, folder, cid, "正文.docx").restored_path)
    assert out.suffix == ".docx"
    assert re.search(r"_恢复自\d{8}-\d{6}\.docx$", out.name), out.name


def test_restore_name_byte_length_within_fs_limit(folder):
    """超长原名:截断主体、绝不截掉时间戳与扩展名,且文件名字节数 ≤255(跨平台安全)。"""
    long_stem = "精算" * 100                          # 200 中文字 ≈ 600 字节,远超 255
    repo, cid = _commit_docx(folder, long_stem + ".docx", "内容")
    out = Path(restore.safe_restore(repo, folder, cid, long_stem + ".docx").restored_path)
    assert len(out.name.encode("utf-8")) <= 255       # 不会被文件系统截断
    assert out.name.endswith(".docx")                 # 扩展名保住
    assert "_恢复自" in out.name                       # 时间戳标记保住
    import re
    assert re.search(r"_恢复自\d{8}-\d{6}\.docx$", out.name)


def test_restore_writes_beside_not_overwrite(folder):
    repo, cid_a = _commit_docx(folder, "正文.docx", "内容A")
    # 当前改成 B 但不提交（模拟未保存改动）
    doc = Document()
    doc.add_paragraph("内容B-未保存")
    doc.save(folder / "正文.docx")
    result = restore.safe_restore(repo, folder, cid_a, "正文.docx")
    # 当前文件未被覆盖
    assert "内容B-未保存" in Document(folder / "正文.docx").paragraphs[0].text
    # 旧版以新文件出现在旁边
    new_file = Path(result.restored_path)
    assert new_file.exists() and new_file != folder / "正文.docx"
    assert "内容A" in Document(new_file).paragraphs[0].text


def test_restore_makes_pre_restore_snapshot(folder):
    repo, cid_a = _commit_docx(folder, "正文.docx", "内容A")
    doc = Document()
    doc.add_paragraph("内容B")
    doc.save(folder / "正文.docx")
    before = len(repo.log())
    restore.safe_restore(repo, folder, cid_a, "正文.docx")
    msgs = [e.message for e in repo.log()]
    assert any("pre-restore" in m for m in msgs)
    assert len(repo.log()) > before


def test_restore_version_skips_deleted_file_in_mixed_version(folder):
    """D2：一版里既改了文件又删了文件——还原这一版只导出还在的文件,
    绝不对被删文件调 show_file(必败、还误报"历史损坏"吓人)。"""
    repo, _ = _commit_docx(folder, "甲.docx", "甲内容")
    _commit_docx(folder, "乙.docx", "乙内容")          # 现有 甲+乙
    # 这一版：删掉甲、改乙(混合版)
    (folder / "甲.docx").unlink()
    doc = Document(); doc.add_paragraph("乙改了"); doc.save(folder / "乙.docx")
    mixed = backup.do_backup(repo, folder).commit_id
    res = restore.restore_version(repo, folder, mixed)
    # 只导出了"乙"这一个还在的文件,没崩
    assert len(res.restored_paths) == 1
    out = Path(res.restored_paths[0])
    assert out.name.startswith("乙") and "恢复自" in out.name
    assert "乙改了" in Document(out).paragraphs[0].text


def test_restore_version_pure_deletion_gives_human_message(folder):
    """D2：这一版的唯一变化就是删文件——没有内容可还原。
    给人话提示(点名被删文件 + 指向上一版),绝不崩、绝不冒"历史损坏"天书。"""
    repo, _ = _commit_docx(folder, "甲.docx", "甲内容")
    _commit_docx(folder, "乙.docx", "乙内容")          # 现有 甲+乙
    (folder / "甲.docx").unlink()                       # 这一版只删甲
    deletion = backup.do_backup(repo, folder).commit_id
    with pytest.raises(DtmError) as e:
        restore.restore_version(repo, folder, deletion)
    assert "甲" in str(e.value)                          # 点名被删文件
    assert "上一版" in str(e.value)                       # 指路
    assert "损坏" not in str(e.value)                     # 不冒吓人的"历史损坏"


def test_restore_version_pre_restore_commit_not_empty_when_clean(folder):
    """D3：工作区干净时 do_backup 不提交、返回空 id;restore 侧须回退到 HEAD,
    不能把空串当 pre_restore_commit 传出去。"""
    repo, cid = _commit_docx(folder, "正文.docx", "内容")   # 提交后工作区干净
    res = restore.restore_version(repo, folder, cid)
    assert res.pre_restore_commit == repo.head()
    assert res.pre_restore_commit                            # 非空


def test_safe_restore_pre_restore_commit_not_empty_when_clean(folder):
    """D3：safe_restore 同样在工作区干净时回退到 HEAD。"""
    repo, cid = _commit_docx(folder, "正文.docx", "内容")
    res = restore.safe_restore(repo, folder, cid, "正文.docx")
    assert res.pre_restore_commit == repo.head()
    assert res.pre_restore_commit


def test_restore_succeeds_despite_stale_office_lock_file(folder):
    # Word 关闭后常残留 ~$ 锁文件;还原是旁存副本(绝不碰原文件,INV-3),不该被它假阻塞。
    # dtm 备份本就把 ~$* 当垃圾忽略,还原也不该把它当"文件开着"的权威信号。
    repo, cid_a = _commit_docx(folder, "正文.docx", "A")
    (folder / "~$正文.docx").write_bytes(b"stale")          # 残留锁文件(Word 已关)
    res = restore.safe_restore(repo, folder, cid_a, "正文.docx")
    assert Path(res.restored_path).exists()                  # 旁存副本照常产出
    assert "恢复自" in Path(res.restored_path).name
    assert (folder / "~$正文.docx").exists()                 # 没去动那个锁文件
